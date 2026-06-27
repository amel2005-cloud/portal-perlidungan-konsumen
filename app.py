import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from io import BytesIO

# ===================== CSS & HERO HEADER =====================
st.markdown("""
    <style>
    .stApp { 
        background-image: radial-gradient(#d1d1d1 0.5px, transparent 0.5px);
        background-size: 20px 20px;
        background-color: #f8f9fa;
    }
    .stTextInput, .stTextArea, .stSelectbox { border-left: 5px solid #b22222; }
    div.stButton > button:first-child { 
        background-color: #b22222; color: white; border: none; font-weight: bold; width: 100%;
    }

    /* ---- HERO BANNER ---- */
    .hero-banner {
        background: linear-gradient(135deg, #8B0000 0%, #b22222 50%, #cc3333 100%);
        border-radius: 12px;
        margin-bottom: 8px;
        position: relative;
        overflow: hidden;
    }
    .hero-banner::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0; bottom: 0;
        background: repeating-linear-gradient(
            45deg, transparent, transparent 20px,
            rgba(255,255,255,0.03) 20px, rgba(255,255,255,0.03) 40px
        );
    }
    .hero-banner::after {
        content: '';
        position: absolute;
        bottom: 0; left: 0; right: 0;
        height: 4px;
        background: linear-gradient(90deg, #FFD700, #FFA500, #FFD700);
    }
    .top-bar {
        background: rgba(0,0,0,0.25);
        padding: 7px 28px;
        display: flex;
        align-items: center;
        justify-content: space-between;
        font-size: 11.5px;
        color: rgba(255,255,255,0.7);
        border-bottom: 1px solid rgba(255,255,255,0.1);
        font-family: Arial, sans-serif;
    }
    .hero-content {
        padding: 28px 32px 32px;
        display: flex;
        align-items: center;
        gap: 28px;
        position: relative;
        z-index: 1;
        flex-wrap: wrap;
    }
    .logo-circle {
        width: 88px; height: 88px;
        border-radius: 50%;
        background: white;
        display: flex; align-items: center; justify-content: center;
        flex-shrink: 0;
        box-shadow: 0 4px 20px rgba(0,0,0,0.3);
        border: 3px solid rgba(255,215,0,0.6);
        font-size: 13px; font-weight: 900; color: #8B0000;
        text-align: center; line-height: 1.3; letter-spacing: 1px;
        font-family: Arial, sans-serif;
    }
    .hero-text { flex: 1; min-width: 200px; }
    .inst-label {
        font-size: 10.5px; color: rgba(255,255,255,0.6);
        letter-spacing: 2.5px; text-transform: uppercase;
        margin-bottom: 6px; font-family: Arial, sans-serif;
    }
    .hero-title {
        font-size: 26px; font-weight: 900; color: white;
        line-height: 1.2; margin-bottom: 4px;
        text-shadow: 0 2px 8px rgba(0,0,0,0.3);
        font-family: Arial, sans-serif;
    }
    .hero-subtitle {
        font-size: 13.5px; color: rgba(255,255,255,0.78);
        margin-bottom: 14px; font-style: italic;
        font-family: Arial, sans-serif;
    }
    .hero-badges { display: flex; gap: 10px; flex-wrap: wrap; }
    .badge {
        background: rgba(255,255,255,0.14);
        border: 1px solid rgba(255,215,0,0.35);
        color: rgba(255,255,255,0.88);
        font-size: 11px; padding: 5px 13px;
        border-radius: 20px;
        display: inline-flex; align-items: center; gap: 6px;
        font-family: Arial, sans-serif;
    }
    .badge-dot { display:inline-block; width:6px; height:6px; border-radius:50%; background:#FFD700; }
    .ojk-badge {
        background: rgba(255,255,255,0.1);
        border: 1px solid rgba(255,215,0,0.45);
        border-radius: 10px; padding: 14px 18px;
        color: white; text-align: right;
        font-family: Arial, sans-serif;
    }
    .ojk-badge-title { font-size: 9.5px; color: rgba(255,255,255,0.55); letter-spacing: 2px; text-transform: uppercase; }
    .ojk-badge-name { font-size: 14px; font-weight: 700; margin-top: 4px; color: #FFD700; }

    /* ---- NAV TABS ---- */
    .nav-tabs {
        background: rgba(0,0,0,0.35);
        display: flex; padding: 0 32px;
        border-top: 1px solid rgba(255,255,255,0.1);
        position: relative; z-index: 1;
        border-bottom: 3px solid #FFD700;
    }
    .nav-tab {
        padding: 11px 18px;
        color: rgba(255,255,255,0.65);
        font-size: 12px; cursor: default;
        border-bottom: 3px solid transparent;
        margin-bottom: -3px;
        letter-spacing: 0.5px;
        font-weight: 700; text-transform: uppercase;
        font-family: Arial, sans-serif;
    }
    .nav-tab.active { color: white; border-bottom-color: #FFD700; background: rgba(255,255,255,0.05); }

    /* ---- INFO RIBBON ---- */
    .info-ribbon {
        background: #fffbf0;
        border-left: 4px solid #FFD700;
        padding: 12px 18px;
        display: flex; align-items: flex-start; gap: 10px;
        border-radius: 0 8px 8px 0;
        font-size: 12.5px; color: #555;
        border: 1px solid #f0e0a0; border-left: 4px solid #FFD700;
        margin-bottom: 4px;
        font-family: Arial, sans-serif; line-height: 1.55;
    }

    /* ---- STAT CARDS ---- */
    .stat-row { display: flex; gap: 12px; margin-bottom: 8px; flex-wrap: wrap; }
    .stat-card {
        background: white; border: 1px solid #e8e8e8;
        border-radius: 8px; padding: 14px 18px;
        flex: 1; min-width: 130px;
        border-top: 3px solid #b22222;
        font-family: Arial, sans-serif;
    }
    .stat-label { font-size: 10px; color: #999; text-transform: uppercase; letter-spacing: 1px; }
    .stat-value { font-size: 20px; font-weight: 700; color: #b22222; margin-top: 4px; }
    .stat-sub { font-size: 10.5px; color: #777; margin-top: 2px; }
    .stat-active { color: #27ae60 !important; }

    /* ---- FORM CARD ---- */
    .form-card-header {
        background: #f8f0f0; border-bottom: 2px solid #b22222;
        padding: 14px 20px; display: flex; align-items: center; gap: 12px;
        border-radius: 8px 8px 0 0;
        font-family: Arial, sans-serif;
    }
    .form-card-icon {
        width: 34px; height: 34px; border-radius: 7px;
        background: #b22222; display: flex; align-items: center;
        justify-content: center; color: white; font-size: 16px;
        flex-shrink: 0;
    }
    .form-card-title { font-size: 13.5px; color: #7a0000; font-weight: 700; }
    .form-card-sub { font-size: 11px; color: #999; margin-top: 2px; }
    </style>

    <!-- HERO BANNER -->
    <div class="hero-banner">
      <div class="top-bar">
        <div>🏛️ Pemerintah Republik Indonesia &nbsp;|&nbsp; 📍 Kantor OJK Jember</div>
        <div>☎ 157 &nbsp;·&nbsp; Hotline Konsumen OJK</div>
      </div>
      <div class="hero-content">
        <div class="logo-circle">OJK<br>JEMBER</div>
        <div class="hero-text">
          <div class="inst-label">Otoritas Jasa Keuangan · Kantor Regional</div>
          <div class="hero-title">Portal Perlindungan Konsumen</div>
          <div class="hero-subtitle">Layanan Surat Resmi Terintegrasi — OJK Jember</div>
          <div class="hero-badges">
            <span class="badge"><span class="badge-dot"></span> Resmi &amp; Terpercaya</span>
            <span class="badge"><span class="badge-dot"></span> Pengaduan Terverifikasi</span>
            <span class="badge"><span class="badge-dot"></span> Layanan Gratis</span>
          </div>
        </div>
        <div class="ojk-badge">
          <div class="ojk-badge-title">Didukung oleh</div>
          <div class="ojk-badge-name">Pemkab Jember</div>
          <div style="font-size:10px;color:rgba(255,255,255,0.5);margin-top:4px;">Layanan Terintegrasi 2026</div>
        </div>
      </div>
      <div class="nav-tabs">
        <div class="nav-tab active">📄 Buat Surat</div>
        <div class="nav-tab">ℹ️ Panduan</div>
        <div class="nav-tab">📞 Kontak</div>
        <div class="nav-tab">❓ FAQ</div>
      </div>
    </div>
""", unsafe_allow_html=True)

# ---- STAT CARDS ----
st.markdown("""
    <div class="stat-row">
      <div class="stat-card">
        <div class="stat-label">Layanan</div>
        <div class="stat-value">2 Jenis</div>
        <div class="stat-sub">Pengaduan &amp; Pernyataan</div>
      </div>
      <div class="stat-card">
        <div class="stat-label">Format</div>
        <div class="stat-value">PDF + DOCX</div>
        <div class="stat-sub">Siap unduh &amp; cetak</div>
      </div>
      <div class="stat-card">
        <div class="stat-label">Status</div>
        <div class="stat-value stat-active">● Aktif</div>
        <div class="stat-sub">Layanan tersedia 24/7</div>
      </div>
    </div>
""", unsafe_allow_html=True)

# ---- INFO RIBBON ----
st.markdown("""
    <div class="info-ribbon">
      ℹ️ &nbsp;<span><strong>Petunjuk Pengisian:</strong> Pastikan NIK sesuai KTP. 
      Isi seluruh kolom dengan data yang benar dan lengkap. 
      Surat yang dihasilkan bersifat resmi dan dapat diajukan langsung ke OJK Jember.</span>
    </div>
""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ---- FORM HEADER ----
st.markdown("""
    <div class="form-card-header">
      <div class="form-card-icon">📋</div>
      <div>
        <div class="form-card-title">FORMULIR LAYANAN SURAT RESMI</div>
        <div class="form-card-sub">Lengkapi data di bawah ini dengan benar untuk membuat surat resmi</div>
      </div>
    </div>
""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ===================== INPUT DATA =====================
kronologis, no_hp, email, pt_dituju = "", "", "", ""
tipe_surat = st.selectbox("Pilih Jenis Layanan:", ["Surat Pengaduan", "Surat Pernyataan"])
nama = st.text_input("Nama Lengkap")
nik = st.text_input("NIK / Nomor Identitas")
if nik and (len(nik) != 16 or not nik.isdigit()):
    st.warning("NIK harus 16 digit angka.")
alamat = st.text_area("Alamat")

if tipe_surat == "Surat Pengaduan":
    pt_dituju = st.text_input("Nama PT yang Dituju")
    no_hp = st.text_input("No. HP")
    email = st.text_input("Email")
    kronologis = st.text_area("Tuliskan kronologis permasalahan:")

kota_ttd = st.text_input("Kota", value="Jember")
tanggal_ttd = st.date_input("Tanggal")

# ===================== FUNGSI PDF & WORD (tidak berubah) =====================
def cetak_baris(pdf, label, nilai):
    pdf.cell(40, 8, txt=label, ln=0)
    pdf.cell(5, 8, txt=":", ln=0)
    pdf.cell(0, 8, txt=nilai, ln=1)

def bersihkan_teks(teks):
    if not teks:
        return ""
    return teks.encode('latin-1', 'replace').decode('latin-1')

TEKS_PERNYATAAN = (
    "menyatakan dengan sesungguhnya bahwa permasalahan yang saya ajukan melalui "
    "Aplikasi Portal Perlindungan Konsumen (APPK) tidak sedang dalam proses atau "
    "pernah diputus oleh lembaga arbitrase atau peradilan atau lembaga mediasi lainnya "
    "termasuk lembaga alternatif penyelesaian Sengketa dan belum pernah difasilitasi oleh OJK."
)
TEKS_PENUTUP = "Demikian surat pernyataan ini dibuat dengan sadar dan tanpa paksaan dari pihak manapun."

def buat_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_font("Arial", size=12)
    if tipe_surat == "Surat Pernyataan":
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, txt="SURAT PERNYATAAN", ln=True, align='C')
        pdf.ln(6)
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 8, txt="Yang bertanda tangan di bawah ini:", ln=True)
        cetak_baris(pdf, "Nama", bersihkan_teks(nama))
        cetak_baris(pdf, "NIK", bersihkan_teks(nik))
        cetak_baris(pdf, "Alamat", bersihkan_teks(alamat))
        pdf.ln(6)
        pdf.multi_cell(0, 8, txt=bersihkan_teks(TEKS_PERNYATAAN))
        pdf.ln(6)
        pdf.multi_cell(0, 8, txt=bersihkan_teks(TEKS_PENUTUP))
    else:
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, txt="SURAT PENGADUAN", ln=True, align='C')
        pdf.ln(6)
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 8, txt=bersihkan_teks(f"Nama PT yang dituju : {pt_dituju}"), ln=True)
        pdf.ln(5)
        pdf.cell(0, 8, txt="Yang bertanda tangan di bawah ini:", ln=True)
        cetak_baris(pdf, "Nama", bersihkan_teks(nama))
        cetak_baris(pdf, "NIK", bersihkan_teks(nik))
        cetak_baris(pdf, "Alamat", bersihkan_teks(alamat))
        cetak_baris(pdf, "No. HP", bersihkan_teks(no_hp))
        cetak_baris(pdf, "Email", bersihkan_teks(email))
        pdf.ln(5)
        pdf.cell(0, 8, txt="Kronologis permasalahan:", ln=True)
        pdf.multi_cell(0, 8, txt=bersihkan_teks(kronologis))
    pdf.ln(20)
    pdf.cell(110)
    pdf.cell(0, 8, txt=bersihkan_teks(f"{kota_ttd}, {tanggal_ttd.strftime('%d %B %Y')}"), ln=True)
    pdf.ln(20)
    pdf.cell(110)
    pdf.cell(0, 8, txt=bersihkan_teks(f"({nama})"), ln=True)
    pdf.output("surat_hasil.pdf")
    with open("surat_hasil.pdf", "rb") as f:
        return f.read()

def buat_word():
    doc = Document()
    if tipe_surat == "Surat Pernyataan":
        judul = doc.add_paragraph("SURAT PERNYATAAN")
        judul.alignment = 1
        judul.runs[0].bold = True
        judul.runs[0].font.size = Pt(14)
        doc.add_paragraph("")
        doc.add_paragraph("Yang bertanda tangan di bawah ini:")
        doc.add_paragraph(f"Nama      : {nama}")
        doc.add_paragraph(f"NIK           : {nik}")
        doc.add_paragraph(f"Alamat    : {alamat}")
        doc.add_paragraph("")
        doc.add_paragraph(TEKS_PERNYATAAN)
        doc.add_paragraph("")
        doc.add_paragraph(TEKS_PENUTUP)
    else:
        judul = doc.add_paragraph("SURAT PENGADUAN")
        judul.alignment = 1
        judul.runs[0].bold = True
        judul.runs[0].font.size = Pt(14)
        doc.add_paragraph("")
        doc.add_paragraph(f"Nama PT yang dituju : {pt_dituju}")
        doc.add_paragraph("")
        doc.add_paragraph("Yang bertanda tangan di bawah ini:")
        doc.add_paragraph(f"Nama      : {nama}")
        doc.add_paragraph(f"NIK           : {nik}")
        doc.add_paragraph(f"Alamat    : {alamat}")
        doc.add_paragraph(f"No. HP   : {no_hp}")
        doc.add_paragraph(f"Email       : {email}")
        doc.add_paragraph("")
        doc.add_paragraph("Kronologis permasalahan:")
        doc.add_paragraph(kronologis)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph(f"                                                  {kota_ttd}, {tanggal_ttd.strftime('%d %B %Y')}")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph(f"                                                  ({nama})")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===================== TOMBOL =====================
if st.button("PROSES & CETAK SURAT"):
    if len(nik) != 16 or not nik.isdigit():
        st.error("⚠️ NIK harus berupa 16 digit angka!")
    else:
        pdf_bytes = buat_pdf()
        word_buffer = buat_word()
        col_pdf, col_word = st.columns(2)
        with col_pdf:
            st.download_button("📄 Download PDF", pdf_bytes, file_name="surat_hasil.pdf")
        with col_word:
            st.download_button("📝 Download Word", word_buffer, file_name="surat_hasil.docx")

st.caption("© 2026 Otoritas Jasa Keuangan & Pemerintah Kabupaten Jember | Layanan Terintegrasi")
